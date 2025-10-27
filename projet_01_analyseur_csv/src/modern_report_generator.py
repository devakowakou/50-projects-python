"""
Module de génération de rapports modernes (PDF, DOCX, PPTX)
Responsabilité: Créer des rapports professionnels avec graphiques intégrés
"""

import pandas as pd
import numpy as np
from datetime import datetime
from typing import Dict, List, Optional
import io
import base64

# PDF
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# DOCX
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PPTX
from pptx import Presentation
from pptx.util import Inches as PptxInches, Pt as PptxPt

# Graphiques
import plotly.graph_objects as go
import plotly.io as pio


class ModernReportGenerator:
    """Générateur de rapports modernes avec graphiques"""
    
    def __init__(self, df: pd.DataFrame, stats_summary: Dict = None, 
                 correlation_data: Dict = None, outliers_data: Dict = None):
        self.df = df
        self.stats_summary = stats_summary or {}
        self.correlation_data = correlation_data or {}
        self.outliers_data = outliers_data or {}
        self.timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        
    def generate_pdf_report(self, filepath: str = None, 
                           include_charts: bool = True,
                           company_name: str = "Analyse de Données") -> str:
        """
        Génère un rapport PDF professionnel
        
        Args:
            filepath: Chemin du fichier
            include_charts: Inclure les graphiques
            company_name: Nom de l'entreprise/projet
            
        Returns:
            Chemin du fichier créé
        """
        if filepath is None:
            filepath = f"rapport_analyse_{self.timestamp}.pdf"
        
        # Créer le document
        doc = SimpleDocTemplate(filepath, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Style personnalisé pour le titre
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        # Style pour les sous-titres
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#1f77b4'),
            spaceAfter=15,
            spaceBefore=15
        )
        
        # Titre (sans emoji pour éviter les problèmes d'encodage)
        story.append(Paragraph("RAPPORT D'ANALYSE DE DONNEES", title_style))
        story.append(Paragraph(f"<b>{company_name}</b>", styles['Heading2']))
        story.append(Paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Résumé Exécutif
        story.append(Paragraph("RESUME EXECUTIF", subtitle_style))
        summary_data = [
            ['Metrique', 'Valeur'],
            ['Nombre de lignes', f"{len(self.df):,}"],
            ['Nombre de colonnes', str(len(self.df.columns))],
            ['Colonnes numeriques', str(len(self.df.select_dtypes(include=['number']).columns))],
            ['Valeurs manquantes', f"{self.df.isnull().sum().sum():,}"],
            ['Taux de completude', f"{(1 - self.df.isnull().sum().sum()/self.df.size)*100:.1f}%"],
        ]
        
        # Largeur adaptative pour éviter la coupure
        t = Table(summary_data, colWidths=[4*inch, 2.5*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(t)
        story.append(Spacer(1, 0.3*inch))
        
        # Statistiques Descriptives
        numeric_cols = self.df.select_dtypes(include=['number']).columns[:5]  # Top 5
        if len(numeric_cols) > 0:
            story.append(Paragraph("STATISTIQUES DESCRIPTIVES", subtitle_style))
            
            for col in numeric_cols:
                # Nettoyage du nom de colonne pour éviter les problèmes d'encodage
                col_name = str(col).encode('ascii', 'ignore').decode('ascii') if isinstance(col, str) else str(col)
                story.append(Paragraph(f"<b>Colonne: {col_name}</b>", styles['Heading3']))
                
                stats_data = [
                    ['Statistique', 'Valeur'],
                    ['Moyenne', f"{self.df[col].mean():.2f}"],
                    ['Mediane', f"{self.df[col].median():.2f}"],
                    ['Ecart-type', f"{self.df[col].std():.2f}"],
                    ['Minimum', f"{self.df[col].min():.2f}"],
                    ['Maximum', f"{self.df[col].max():.2f}"],
                ]
                
                # Largeur adaptée pour éviter la coupure
                t = Table(stats_data, colWidths=[3*inch, 2.5*inch])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                story.append(t)
                story.append(Spacer(1, 0.2*inch))
        
        # Recommandations
        story.append(PageBreak())
        story.append(Paragraph("RECOMMANDATIONS", subtitle_style))
        
        recommendations = self._generate_recommendations()
        for i, rec in enumerate(recommendations, 1):
            # Nettoyer la recommandation des emojis
            clean_rec = rec.encode('ascii', 'ignore').decode('ascii')
            story.append(Paragraph(f"<b>{i}.</b> {clean_rec}", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        
        # Générer le PDF
        doc.build(story)
        return filepath
    
    def generate_docx_report(self, filepath: str = None,
                            company_name: str = "Analyse de Données") -> str:
        """
        Génère un rapport DOCX (Word) professionnel
        
        Args:
            filepath: Chemin du fichier
            company_name: Nom de l'entreprise/projet
            
        Returns:
            Chemin du fichier créé
        """
        if filepath is None:
            filepath = f"rapport_analyse_{self.timestamp}.docx"
        
        doc = Document()
        
        # Titre principal (sans emoji pour meilleure compatibilité)
        title = doc.add_heading('RAPPORT D\'ANALYSE DE DONNEES', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sous-titre
        subtitle = doc.add_heading(company_name, level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espace
        
        # Résumé Exécutif (sans emoji)
        doc.add_heading('RESUME EXECUTIF', level=1)
        
        # Table résumé avec largeurs adaptées
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Définir les largeurs de colonnes pour éviter la coupure
        for row in table.rows:
            row.cells[0].width = Inches(3.5)
            row.cells[1].width = Inches(2.5)
        
        summary_data = [
            ('Metrique', 'Valeur'),
            ('Nombre de lignes', f"{len(self.df):,}"),
            ('Nombre de colonnes', str(len(self.df.columns))),
            ('Colonnes numeriques', str(len(self.df.select_dtypes(include=['number']).columns))),
            ('Valeurs manquantes', f"{self.df.isnull().sum().sum():,}"),
            ('Taux de completude', f"{(1 - self.df.isnull().sum().sum()/self.df.size)*100:.1f}%"),
            ('Lignes dupliquees', str(self.df.duplicated().sum())),
        ]
        
        for i, (key, value) in enumerate(summary_data):
            row = table.rows[i]
            row.cells[0].text = key
            row.cells[1].text = value
            
            # Style de la première ligne (en-tête)
            if i == 0:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Statistiques Descriptives (sans emoji)
        doc.add_heading('STATISTIQUES DESCRIPTIVES', level=1)
        
        numeric_cols = self.df.select_dtypes(include=['number']).columns[:5]
        for col in numeric_cols:
            # Nettoyer le nom de colonne
            col_name = str(col).encode('ascii', 'ignore').decode('ascii') if isinstance(col, str) else str(col)
            doc.add_heading(f'Colonne: {col_name}', level=2)
            
            stats_table = doc.add_table(rows=6, cols=2)
            stats_table.style = 'Light List Accent 1'
            
            # Définir largeurs pour éviter la coupure
            for row in stats_table.rows:
                row.cells[0].width = Inches(2.5)
                row.cells[1].width = Inches(2.5)
            
            stats = [
                ('Moyenne', f"{self.df[col].mean():.2f}"),
                ('Mediane', f"{self.df[col].median():.2f}"),
                ('Ecart-type', f"{self.df[col].std():.2f}"),
                ('Minimum', f"{self.df[col].min():.2f}"),
                ('Maximum', f"{self.df[col].max():.2f}"),
            ]
            
            for i, (stat_name, stat_value) in enumerate(stats):
                row = stats_table.rows[i]
                row.cells[0].text = stat_name
                row.cells[1].text = stat_value
        
        # Recommandations
        doc.add_page_break()
        doc.add_heading('RECOMMANDATIONS', level=1)
        
        recommendations = self._generate_recommendations()
        for i, rec in enumerate(recommendations, 1):
            # Nettoyer la recommandation des emojis
            clean_rec = rec.encode('ascii', 'ignore').decode('ascii')
            p = doc.add_paragraph(clean_rec, style='List Number')
        
        # Conclusion
        doc.add_paragraph()
        doc.add_heading('CONCLUSION', level=1)
        conclusion = doc.add_paragraph(
            f"Ce rapport a été généré automatiquement le {datetime.now().strftime('%d/%m/%Y à %H:%M')}. "
            f"L'analyse porte sur un dataset de {len(self.df):,} lignes et {len(self.df.columns)} colonnes."
        )
        
        # Sauvegarder
        doc.save(filepath)
        return filepath
    
    def generate_html_report(self, filepath: str = None, 
                            include_interactive_charts: bool = True) -> str:
        """
        Génère un rapport HTML interactif
        
        Args:
            filepath: Chemin du fichier
            include_interactive_charts: Inclure graphiques interactifs Plotly
            
        Returns:
            Chemin du fichier créé
        """
        if filepath is None:
            filepath = f"rapport_analyse_{self.timestamp}.html"
        
        html_content = f"""
<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Rapport d'Analyse de Données</title>
    <script src="https://cdn.plot.ly/plotly-latest.min.js"></script>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        }}
        .container {{
            background: white;
            padding: 40px;
            border-radius: 15px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        }}
        h1 {{
            color: #1f77b4;
            text-align: center;
            border-bottom: 3px solid #1f77b4;
            padding-bottom: 20px;
        }}
        h2 {{
            color: #333;
            margin-top: 30px;
            border-left: 5px solid #1f77b4;
            padding-left: 15px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        th {{
            background: #1f77b4;
            color: white;
            padding: 12px;
            text-align: left;
        }}
        td {{
            padding: 10px;
            border-bottom: 1px solid #ddd;
        }}
        tr:hover {{
            background: #f5f5f5;
        }}
        .metric {{
            display: inline-block;
            background: #f8f9fa;
            padding: 15px 25px;
            margin: 10px;
            border-radius: 8px;
            border-left: 4px solid #1f77b4;
        }}
        .metric-value {{
            font-size: 24px;
            font-weight: bold;
            color: #1f77b4;
        }}
        .metric-label {{
            font-size: 12px;
            color: #666;
            text-transform: uppercase;
        }}
        .recommendation {{
            background: #e7f3ff;
            padding: 15px;
            margin: 10px 0;
            border-left: 4px solid #0066cc;
            border-radius: 5px;
        }}
        .footer {{
            text-align: center;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            color: #666;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>📊 Rapport d'Analyse de Données</h1>
        <p style="text-align: center; color: #666;">
            Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}
        </p>
        
        <h2>📈 Résumé Exécutif</h2>
        <div style="text-align: center;">
            <div class="metric">
                <div class="metric-value">{len(self.df):,}</div>
                <div class="metric-label">Lignes</div>
            </div>
            <div class="metric">
                <div class="metric-value">{len(self.df.columns)}</div>
                <div class="metric-label">Colonnes</div>
            </div>
            <div class="metric">
                <div class="metric-value">{(1 - self.df.isnull().sum().sum()/self.df.size)*100:.1f}%</div>
                <div class="metric-label">Complétude</div>
            </div>
            <div class="metric">
                <div class="metric-value">{self.df.duplicated().sum()}</div>
                <div class="metric-label">Duplicatas</div>
            </div>
        </div>
        
        <h2>📊 Statistiques Descriptives</h2>
        {self._generate_stats_html_table()}
        
        <h2>💡 Recommandations</h2>
        {self._generate_recommendations_html()}
        
        <div class="footer">
            <p>Rapport généré automatiquement par Analyseur CSV Professionnel</p>
            <p>© 2025 - Fait avec ❤️ et Python</p>
        </div>
    </div>
</body>
</html>
"""
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return filepath
    
    def _generate_stats_html_table(self) -> str:
        """Génère le HTML pour les statistiques"""
        numeric_cols = self.df.select_dtypes(include=['number']).columns[:5]
        
        if len(numeric_cols) == 0:
            return "<p>Aucune colonne numérique à analyser.</p>"
        
        html = "<table><thead><tr><th>Colonne</th><th>Moyenne</th><th>Mediane</th><th>Ecart-type</th><th>Min</th><th>Max</th></tr></thead><tbody>"
        
        for col in numeric_cols:
            html += f"""
            <tr>
                <td><strong>{col}</strong></td>
                <td>{self.df[col].mean():.2f}</td>
                <td>{self.df[col].median():.2f}</td>
                <td>{self.df[col].std():.2f}</td>
                <td>{self.df[col].min():.2f}</td>
                <td>{self.df[col].max():.2f}</td>
            </tr>
            """
        
        html += "</tbody></table>"
        return html
    
    def _generate_recommendations(self) -> List[str]:
        """Génère des recommandations automatiques basées sur l'analyse"""
        recommendations = []
        
        # Vérifier les valeurs manquantes
        missing_pct = (self.df.isnull().sum().sum() / self.df.size) * 100
        if missing_pct > 5:
            recommendations.append(
                f"ATTENTION: Le dataset contient {missing_pct:.1f}% de valeurs manquantes. "
                "Envisagez un nettoyage des donnees ou une imputation appropriee."
            )
        elif missing_pct > 0:
            recommendations.append(
                f"Le dataset contient {missing_pct:.1f}% de valeurs manquantes, ce qui est acceptable."
            )
        else:
            recommendations.append("Excellente qualite : aucune valeur manquante detectee.")
        
        # Vérifier les duplicatas
        duplicates = self.df.duplicated().sum()
        if duplicates > 0:
            recommendations.append(
                f"INFO: {duplicates} lignes dupliquees detectees ({duplicates/len(self.df)*100:.1f}%). "
                "Considerez leur suppression si non intentionnelles."
            )
        else:
            recommendations.append("Aucune ligne dupliquee detectee.")
        
        # Analyser les colonnes numériques
        numeric_cols = self.df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            recommendations.append(
                f"STATISTIQUES: {len(numeric_cols)} colonnes numeriques disponibles pour des analyses statistiques avancees."
            )
            
            # Vérifier la variance
            for col in numeric_cols:
                cv = (self.df[col].std() / self.df[col].mean()) * 100 if self.df[col].mean() != 0 else 0
                if cv > 100:
                    col_name = str(col).encode('ascii', 'ignore').decode('ascii') if isinstance(col, str) else str(col)
                    recommendations.append(
                        f"VARIABILITE: La colonne '{col_name}' a une forte variabilite (CV = {cv:.1f}%). "
                        "Considerez une normalisation pour certaines analyses."
                    )
        
        # Taille du dataset
        if len(self.df) < 100:
            recommendations.append(
                "ATTENTION: Dataset de petite taille. Les analyses statistiques peuvent etre moins fiables."
            )
        elif len(self.df) > 10000:
            recommendations.append(
                "Dataset de grande taille excellent pour des analyses robustes."
            )
        
        return recommendations
    
    def _generate_recommendations_html(self) -> str:
        """Génère le HTML pour les recommandations"""
        recommendations = self._generate_recommendations()
        html = ""
        
        for i, rec in enumerate(recommendations, 1):
            html += f'<div class="recommendation">{i}. {rec}</div>'
        
        return html
